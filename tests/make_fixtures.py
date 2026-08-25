"""
Prepare fixtures for tests/selftest.js.

Extracts the REAL injected JavaScript (same code that ships) and builds a
sample article from a generated .docx, so the JS test runs against
production strings rather than copies.

Deliberately avoids importing vcpaste.app, because that imports webview,
which is Windows-only: the top-bar JS is pulled from the source via ast.

Usage:  python tests/make_fixtures.py
"""
import ast
import io
import json
import os
import sys

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, ROOT)
OUT = os.path.join(ROOT, "tests", "build")

from vcpaste import config, converter, inject_js  # noqa: E402


def top_bar_js() -> str:
    """Pull _build_top_bar_js's return value without importing webview."""
    src = open(os.path.join(ROOT, "vcpaste", "app.py"), encoding="utf-8").read()
    fn = next(n for n in ast.parse(src).body
              if isinstance(n, ast.FunctionDef) and n.name == "_build_top_bar_js")
    raw = None
    for node in ast.walk(fn):
        if isinstance(node, ast.Return):
            v = node.value
            if isinstance(v, ast.Constant):
                raw = v.value
            elif isinstance(v, ast.Call):          # "...".replace().replace()
                inner = v
                while isinstance(inner, ast.Call):
                    inner = inner.func.value
                if isinstance(inner, ast.Constant):
                    raw = inner.value
    if raw is None:
        raise SystemExit("не удалось извлечь JS верхней панели из app.py")
    return (raw.replace("__EDITOR__", json.dumps(config.EDITOR_URL))
               .replace("__LOGIN__", json.dumps(config.LOGIN_URL)))


def sample_article() -> dict:
    """A .docx with a heading, text and an image -> converted HTML."""
    from docx import Document
    from PIL import Image

    png = io.BytesIO()
    Image.new("RGB", (60, 40), (18, 184, 105)).save(png, "PNG")
    doc = Document()
    doc.add_heading("Тестовая статья", 0)
    doc.add_paragraph("Первый абзац с текстом.")
    doc.add_picture(io.BytesIO(png.getvalue()))
    buf = io.BytesIO()
    doc.save(buf)

    data = converter.convert_bytes("Тестовая статья.docx", buf.getvalue())
    assert data["images"] == 1, "картинка потерялась при конвертации"
    assert "data:image/png;base64" in data["html"], "картинка не встроена"
    return data


def main():
    os.makedirs(OUT, exist_ok=True)
    files = {
        "insert.js": inject_js.build_insert_js(
            config.EDITOR_AREA_SELECTORS, config.EDITOR_TITLE_SELECTORS),
        "bar.js": top_bar_js(),
        "autofill.js": inject_js.build_autofill_js(
            "user@example.com", "SavedPass123",
            config.LOGIN_EMAIL_SELECTORS, config.LOGIN_PASSWORD_SELECTORS,
            config.LOGIN_SUBMIT_SELECTORS, config.LOGIN_SUBMIT_TEXTS),
        "article.json": json.dumps(sample_article(), ensure_ascii=False),
    }
    for name, text in files.items():
        with open(os.path.join(OUT, name), "w", encoding="utf-8") as f:
            f.write(text)
        print("написан", name, len(text), "симв.")

    # Python-side sanity checks that need no browser.
    for bad in ("x.pdf", "x.rtf"):
        try:
            converter.convert_bytes(bad, b"junk")
        except ValueError:
            pass
        else:
            raise SystemExit(f"{bad} должен отклоняться")
    print("проверки конвертера пройдены")


if __name__ == "__main__":
    main()
